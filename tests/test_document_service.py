"""Document AI extraction tests."""
from __future__ import annotations

import io
import zipfile

import pytest
from docx import Document as WordDocument

from document_service import (
    DocumentServiceError,
    PdfExtraction,
    build_document_request,
    build_pdf_request,
    extract_docx_text,
    extract_pdf_text,
    extract_plain_text,
)


def _docx_bytes() -> bytes:
    document = WordDocument()
    document.add_paragraph("Mahnung vom Jobcenter")
    document.add_paragraph("Bitte reagieren Sie bis 10.08.2026.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Betrag"
    table.cell(0, 1).text = "40 Euro"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_empty_and_invalid_pdf_are_rejected() -> None:
    with pytest.raises(DocumentServiceError, match="empty"):
        extract_pdf_text(b"")
    with pytest.raises(DocumentServiceError, match="invalid"):
        extract_pdf_text(b"not-a-pdf")


def test_blank_pdf_is_classified_as_scanned(monkeypatch) -> None:
    class BlankPage:
        def extract_text(self) -> str:
            return ""

    class Reader:
        is_encrypted = False
        pages = [BlankPage()]

    monkeypatch.setattr("document_service.PdfReader", lambda *_args, **_kwargs: Reader())
    with pytest.raises(DocumentServiceError, match="scanned"):
        extract_pdf_text(b"%PDF-blank")


def test_pdf_extraction_is_clean_bounded_and_reports_metadata(monkeypatch) -> None:
    class Page:
        def extract_text(self) -> str:
            return "  Mahnung   vom Jobcenter\n\n\nFrist: 10.08.2026. Bitte reagieren Sie rechtzeitig.  "

    class Reader:
        is_encrypted = False
        pages = [Page(), Page()]

    monkeypatch.setattr("document_service.PdfReader", lambda *_args, **_kwargs: Reader())
    result = extract_pdf_text(b"%PDF-test")

    assert result.page_count == 2
    assert result.pages_read == 2
    assert "Mahnung vom Jobcenter" in result.text
    assert "Frist: 10.08.2026" in result.text
    assert result.truncated is False


def test_docx_extracts_paragraphs_and_table_cells() -> None:
    result = extract_docx_text(_docx_bytes())
    assert result.kind == "docx"
    assert "Mahnung vom Jobcenter" in result.text
    assert "Betrag | 40 Euro" in result.text
    assert result.truncated is False


def test_plain_text_supports_utf8_and_cp1252() -> None:
    utf8 = extract_plain_text("Frist: 10.08.2026. Bitte antworten.".encode("utf-8"))
    cp1252 = extract_plain_text("Kündigung spätestens im März möglich.".encode("cp1252"))
    assert utf8.kind == "text" and "10.08.2026" in utf8.text
    assert "März" in cp1252.text


def test_docx_archive_with_excessive_file_count_is_rejected() -> None:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("word/document.xml", "<xml/>")
        for index in range(501):
            archive.writestr(f"word/media/{index}.txt", "x")
    with pytest.raises(DocumentServiceError, match="unsafe_archive"):
        extract_docx_text(buffer.getvalue())


def test_document_requests_use_user_language_and_document_text() -> None:
    extraction = PdfExtraction(
        text="Rechnung 123. Frist 10.08.2026.",
        page_count=3,
        pages_read=3,
        truncated=False,
    )
    pdf_prompt = build_pdf_request(extraction, language="ar", note="rechnung.pdf")
    text_prompt = build_document_request(
        extract_plain_text(b"Invoice text with enough useful characters for extraction."),
        language="en",
        note="notice.txt",
    )

    assert "بالعربية فقط" in pdf_prompt
    assert "rechnung.pdf" in pdf_prompt
    assert "Rechnung 123" in pdf_prompt
    assert "Document kind: pdf" in pdf_prompt
    assert "Please explain this document in English" in text_prompt
    assert "notice.txt" in text_prompt
