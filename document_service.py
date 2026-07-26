"""Privacy-conscious document extraction for AmtHero24."""
from __future__ import annotations

import io
import re
import unicodedata
import zipfile
from dataclasses import dataclass

from docx import Document
from pypdf import PdfReader

_MAX_DOCUMENT_BYTES = 20 * 1024 * 1024
_MAX_DOCX_UNCOMPRESSED_BYTES = 40 * 1024 * 1024
_MAX_DOCX_FILES = 500
_MAX_PAGES = 30
_MAX_TEXT_CHARACTERS = 16_000


class DocumentServiceError(RuntimeError):
    """Document processing failure with a stable user-facing category."""

    def __init__(self, code: str) -> None:
        super().__init__(code)
        self.code = code


@dataclass(frozen=True)
class DocumentExtraction:
    text: str
    kind: str
    units_total: int
    units_read: int
    truncated: bool


@dataclass(frozen=True)
class PdfExtraction:
    """Backward-compatible PDF result used by existing tests and callers."""

    text: str
    page_count: int
    pages_read: int
    truncated: bool


def _clean_text(value: str) -> str:
    text = unicodedata.normalize("NFKC", value or "")
    text = "".join(
        character
        for character in text
        if unicodedata.category(character) not in {"Cf", "Cs"} or character in {"\n", "\t"}
    )
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _validate_bytes(content: bytes) -> None:
    if not content:
        raise DocumentServiceError("empty")
    if len(content) > _MAX_DOCUMENT_BYTES:
        raise DocumentServiceError("too_large")


def _bounded_join(chunks: list[str]) -> tuple[str, bool]:
    output: list[str] = []
    used = 0
    truncated = False
    for chunk in chunks:
        cleaned = _clean_text(chunk)
        if not cleaned:
            continue
        remaining = _MAX_TEXT_CHARACTERS - used
        if remaining <= 0:
            truncated = True
            break
        output.append(cleaned[:remaining])
        used += len(output[-1])
        if len(cleaned) > remaining:
            truncated = True
            break
    return _clean_text("\n\n".join(output)), truncated


def extract_pdf_text(pdf_bytes: bytes) -> PdfExtraction:
    """Extract bounded text from a text-based PDF without storing the file."""
    _validate_bytes(pdf_bytes)
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes), strict=False)
    except Exception as exc:
        raise DocumentServiceError("invalid") from exc

    if reader.is_encrypted:
        try:
            unlocked = reader.decrypt("")
        except Exception as exc:
            raise DocumentServiceError("encrypted") from exc
        if not unlocked:
            raise DocumentServiceError("encrypted")

    page_count = len(reader.pages)
    pages_read = min(page_count, _MAX_PAGES)
    chunks: list[str] = []
    for page in reader.pages[:pages_read]:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:
            chunks.append("")
    text, text_truncated = _bounded_join(chunks)
    if len(text) < 40:
        raise DocumentServiceError("scanned")
    return PdfExtraction(
        text=text,
        page_count=page_count,
        pages_read=pages_read,
        truncated=page_count > _MAX_PAGES or text_truncated,
    )


def _validate_docx_archive(docx_bytes: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as archive:
            members = archive.infolist()
            if len(members) > _MAX_DOCX_FILES:
                raise DocumentServiceError("unsafe_archive")
            total_size = sum(max(0, member.file_size) for member in members)
            if total_size > _MAX_DOCX_UNCOMPRESSED_BYTES:
                raise DocumentServiceError("unsafe_archive")
            if not any(member.filename == "word/document.xml" for member in members):
                raise DocumentServiceError("invalid")
    except DocumentServiceError:
        raise
    except (zipfile.BadZipFile, OSError) as exc:
        raise DocumentServiceError("invalid") from exc


def extract_docx_text(docx_bytes: bytes) -> DocumentExtraction:
    """Extract paragraphs and table cells from a modern Word document."""
    _validate_bytes(docx_bytes)
    _validate_docx_archive(docx_bytes)
    try:
        document = Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        raise DocumentServiceError("invalid") from exc

    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [_clean_text(cell.text) for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                chunks.append(row_text)

    text, truncated = _bounded_join(chunks)
    if len(text) < 20:
        raise DocumentServiceError("empty_text")
    return DocumentExtraction(
        text=text,
        kind="docx",
        units_total=len(document.paragraphs) + len(document.tables),
        units_read=len(document.paragraphs) + len(document.tables),
        truncated=truncated,
    )


def extract_plain_text(text_bytes: bytes) -> DocumentExtraction:
    """Decode a bounded plain-text document with conservative fallbacks."""
    _validate_bytes(text_bytes)
    decoded = ""
    for encoding in ("utf-8-sig", "utf-16", "cp1252"):
        try:
            decoded = text_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if not decoded:
        raise DocumentServiceError("invalid_encoding")
    text, truncated = _bounded_join([decoded])
    if len(text) < 20:
        raise DocumentServiceError("empty_text")
    return DocumentExtraction(
        text=text,
        kind="text",
        units_total=max(1, decoded.count("\n") + 1),
        units_read=max(1, decoded.count("\n") + 1),
        truncated=truncated,
    )


def _instruction(language: str) -> str:
    headers = {
        "ar": "اشرح هذا المستند بالعربية فقط. أعطني المعنى الأساسي، أهم موعد أو مبلغ إن وُجد، والخطوة العملية التالية.",
        "de": "Erkläre dieses Dokument auf Deutsch. Nenne die Kernaussage, wichtige Fristen oder Beträge und den nächsten praktischen Schritt.",
        "en": "Please explain this document in English. Give the main meaning, any important deadline or amount, and the next practical step.",
        "uk": "Поясни цей документ українською. Назви головний зміст, важливий термін або суму та наступний практичний крок.",
        "el": "Εξήγησε αυτό το έγγραφο στα Ελληνικά. Δώσε το βασικό νόημα, σημαντική προθεσμία ή ποσό και το επόμενο πρακτικό βήμα.",
    }
    return headers.get(language, headers["de"])


def build_document_request(extraction: DocumentExtraction, *, language: str, note: str = "") -> str:
    safe_note = " ".join((note or "").split())[:180]
    metadata = (
        f"Document kind: {extraction.kind}; units: {extraction.units_total}; "
        f"units read: {extraction.units_read}; truncated: {str(extraction.truncated).lower()}."
    )
    note_line = f"User note or filename: {safe_note}\n" if safe_note else ""
    return f"{_instruction(language)}\n{metadata}\n{note_line}Extracted document text:\n{extraction.text}"


def build_pdf_request(extraction: PdfExtraction, *, language: str, note: str = "") -> str:
    """Backward-compatible wrapper for PDF prompt construction."""
    generic = DocumentExtraction(
        text=extraction.text,
        kind="pdf",
        units_total=extraction.page_count,
        units_read=extraction.pages_read,
        truncated=extraction.truncated,
    )
    return build_document_request(generic, language=language, note=note)
